#!/usr/bin/env python3
"""Rebuild the Word template for research notes with enriched sections.

Adds: Five Key Tensions, News & Events, Cost Structure & Margin Drivers,
Industry Deep Dive, What You Need to Believe (replaces Scenario Analysis),
Forward Catalysts, and Monitoring Framework.

Usage:
    python scripts/update_research_template.py

Output:
    templates/research_note.docx
"""

import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
DARK_GRAY = RGBColor(0x4A, 0x4A, 0x4A)
MEDIUM_GRAY = RGBColor(0x6B, 0x6B, 0x6B)

FONT_NAME = "Calibri"
TITLE_SIZE = Pt(28)
HEADING1_SIZE = Pt(16)
HEADING2_SIZE = Pt(13)
NORMAL_SIZE = Pt(11)

OUTPUT_PATH = os.path.join("templates", "research_note.docx")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _configure_styles(doc):
    """Set up document styles for Title, Heading 1, Heading 2, and Normal."""
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = NORMAL_SIZE
    normal.font.color.rgb = DARK_GRAY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    h1 = styles["Heading 1"]
    h1.font.name = FONT_NAME
    h1.font.size = HEADING1_SIZE
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)

    h2 = styles["Heading 2"]
    h2.font.name = FONT_NAME
    h2.font.size = HEADING2_SIZE
    h2.font.bold = True
    h2.font.color.rgb = DARK_GRAY
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    title = styles["Title"]
    title.font.name = FONT_NAME
    title.font.size = TITLE_SIZE
    title.font.color.rgb = NAVY
    title.font.bold = False
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)

    subtitle = styles["Subtitle"]
    subtitle.font.name = FONT_NAME
    subtitle.font.size = Pt(16)
    subtitle.font.color.rgb = DARK_GRAY
    subtitle.font.bold = False
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)


def _set_page_margins(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def _add_page_break(doc):
    doc.add_page_break()


def _add_centered_paragraph(doc, text, style="Normal", color=None):
    para = doc.add_paragraph(text, style=style)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if color:
        for run in para.runs:
            run.font.color.rgb = color
    return para


# ---------------------------------------------------------------------------
# Template sections
# ---------------------------------------------------------------------------

def _add_cover_page(doc):
    for _ in range(6):
        doc.add_paragraph("")
    doc.add_paragraph("{{ company_name }} ({{ ticker }})", style="Title")
    doc.add_paragraph("Research Note \u2014 {{ date }}", style="Subtitle")
    _add_centered_paragraph(
        doc, "Price: {{ price }} | Market Cap: {{ market_cap }}",
        color=MEDIUM_GRAY,
    )


def _add_executive_summary(doc):
    _add_page_break(doc)
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("{{ five_key_tensions }}")
    doc.add_paragraph("{{ executive_summary }}")
    doc.add_paragraph("{{p key_metrics_table }}")


def _add_investment_thesis(doc):
    doc.add_heading("Investment Thesis", level=1)
    doc.add_paragraph("{{ investment_thesis }}")
    doc.add_heading("Variant Perception", level=2)
    doc.add_paragraph("{{ variant_perception }}")


def _add_company_overview(doc):
    doc.add_heading("Company Overview", level=1)
    doc.add_paragraph("{{ company_description }}")


def _add_news_and_events(doc):
    doc.add_heading("Key News & Events", level=1)
    doc.add_paragraph("{{ news_timeline }}")


def _add_financial_analysis(doc):
    doc.add_heading("Financial Analysis", level=1)
    doc.add_paragraph("{{ revenue_chart }}")
    doc.add_paragraph("{{p financials_table }}")
    doc.add_paragraph("{{ margin_chart }}")

    doc.add_heading("Cost Structure & Margin Drivers", level=2)
    doc.add_paragraph("{{ cost_margin_analysis }}")
    doc.add_paragraph("{{p opex_breakdown_table }}")

    doc.add_heading("Revenue by Segment", level=2)
    doc.add_paragraph("{{ segment_chart }}")
    doc.add_paragraph("{{p segments_table }}")

    doc.add_heading("Revenue by Geography", level=2)
    doc.add_paragraph("{{p geo_table }}")

    doc.add_heading("Capital Return & Share Count", level=2)
    doc.add_paragraph("{{p shares_outstanding_table }}")


def _add_industry_deep_dive(doc):
    doc.add_heading("Industry Deep Dive", level=1)
    doc.add_paragraph("{{ industry_deep_dive }}")


def _add_guidance_track_record(doc):
    doc.add_paragraph("{% if has_guidance %}")
    doc.add_heading("Guidance Track Record", level=1)
    doc.add_paragraph("{{ guidance_track_record }}")
    doc.add_paragraph("{% endif %}")


def _add_what_you_need_to_believe(doc):
    doc.add_heading("What You Need to Believe", level=1)

    doc.add_heading("To Go Long", level=2)
    doc.add_paragraph("{{ bull_beliefs }}")
    doc.add_paragraph("Price Target: {{ bull_target }}")

    doc.add_heading("To Go Short", level=2)
    doc.add_paragraph("{{ bear_beliefs }}")
    doc.add_paragraph("Price Target: {{ bear_target }}")

    doc.add_heading("Risk/Reward Assessment", level=2)
    doc.add_paragraph("{{ risk_reward_assessment }}")


def _add_forward_catalysts(doc):
    doc.add_heading("Forward Catalysts", level=1)
    doc.add_paragraph("{{ forward_catalysts }}")
    doc.add_paragraph("{{ policy_backdrop }}")


def _add_capital_allocation(doc):
    doc.add_heading("Capital Allocation", level=1)
    doc.add_paragraph("{{ capital_allocation_commentary }}")


def _add_valuation(doc):
    doc.add_heading("Valuation", level=1)

    doc.add_paragraph("{% if has_dcf %}")
    doc.add_heading("DCF Analysis", level=2)
    doc.add_paragraph("{{ dcf_summary }}")
    doc.add_paragraph("{{ dcf_sensitivity_chart }}")
    doc.add_paragraph("{% endif %}")

    doc.add_paragraph("{% if has_comps %}")
    doc.add_heading("Comparable Companies", level=2)
    doc.add_paragraph("{{ comps_commentary }}")
    doc.add_paragraph("{% endif %}")


def _add_risks(doc):
    doc.add_heading("Key Risks", level=1)
    doc.add_paragraph("{{ risks_summary }}")


def _add_monitoring_framework(doc):
    doc.add_heading("Monitoring Framework", level=1)
    doc.add_heading("Quantitative Monitors", level=2)
    doc.add_paragraph("{{ monitoring_quantitative }}")
    doc.add_heading("Qualitative Monitors", level=2)
    doc.add_paragraph("{{ monitoring_qualitative }}")


def _add_appendix(doc):
    doc.add_heading("Appendix", level=1)
    doc.add_paragraph("{{ appendix_content }}")
    doc.add_paragraph("")
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("Data sourced from Daloopa")
    run.font.size = Pt(9)
    run.font.color.rgb = MEDIUM_GRAY
    run.font.italic = True


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    doc = Document()
    _configure_styles(doc)
    _set_page_margins(doc)

    _add_cover_page(doc)
    _add_executive_summary(doc)
    _add_investment_thesis(doc)
    _add_company_overview(doc)
    _add_news_and_events(doc)
    _add_financial_analysis(doc)
    _add_industry_deep_dive(doc)
    _add_guidance_track_record(doc)
    _add_what_you_need_to_believe(doc)
    _add_forward_catalysts(doc)
    _add_capital_allocation(doc)
    _add_valuation(doc)
    _add_risks(doc)
    _add_monitoring_framework(doc)
    _add_appendix(doc)

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
