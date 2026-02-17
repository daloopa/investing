#!/usr/bin/env python3
"""Create the Word template for research notes using python-docx.

Usage:
    python scripts/create_template.py

Output:
    templates/research_note.docx

The generated .docx contains Jinja2 template tags that docxtpl will
process at render time. Tables are built dynamically by the renderer
using Subdoc objects — this template only has paragraph placeholders
for them (e.g. {{ key_metrics_table }}).
"""

import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
DARK_GRAY = RGBColor(0x4A, 0x4A, 0x4A)
MEDIUM_GRAY = RGBColor(0x6B, 0x6B, 0x6B)

FONT_NAME = "Calibri"
TITLE_SIZE = Pt(28)
HEADING1_SIZE = Pt(16)
HEADING2_SIZE = Pt(13)
NORMAL_SIZE = Pt(11)

OUTPUT_PATH = os.path.join("templates", "research_note.docx")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _configure_styles(doc):
    """Set up document styles for Title, Heading 1, Heading 2, and Normal."""
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = NORMAL_SIZE
    normal.font.color.rgb = DARK_GRAY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    h1 = styles["Heading 1"]
    h1.font.name = FONT_NAME
    h1.font.size = HEADING1_SIZE
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)

    h2 = styles["Heading 2"]
    h2.font.name = FONT_NAME
    h2.font.size = HEADING2_SIZE
    h2.font.bold = True
    h2.font.color.rgb = DARK_GRAY
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    title = styles["Title"]
    title.font.name = FONT_NAME
    title.font.size = TITLE_SIZE
    title.font.color.rgb = NAVY
    title.font.bold = False
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)

    subtitle = styles["Subtitle"]
    subtitle.font.name = FONT_NAME
    subtitle.font.size = Pt(16)
    subtitle.font.color.rgb = DARK_GRAY
    subtitle.font.bold = False
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)


def _set_page_margins(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def _add_page_break(doc):
    doc.add_page_break()


def _add_centered_paragraph(doc, text, style="Normal", color=None):
    para = doc.add_paragraph(text, style=style)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if color:
        for run in para.runs:
            run.font.color.rgb = color
    return para


# ---------------------------------------------------------------------------
# Template sections
# ---------------------------------------------------------------------------

def _add_cover_page(doc):
    for _ in range(6):
        doc.add_paragraph("")
    doc.add_paragraph("{{ company_name }} ({{ ticker }})", style="Title")
    doc.add_paragraph("Research Note \u2014 {{ date }}", style="Subtitle")
    _add_centered_paragraph(
        doc, "Price: {{ price }} | Market Cap: {{ market_cap }}",
        color=MEDIUM_GRAY,
    )


def _add_executive_summary(doc):
    _add_page_break(doc)
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("{{ executive_summary }}")
    # Table built dynamically by renderer as Subdoc
    doc.add_paragraph("{{ key_metrics_table }}")


def _add_investment_thesis(doc):
    doc.add_heading("Investment Thesis", level=1)
    doc.add_paragraph("{{ variant_perception }}")
    doc.add_heading("Thesis Pillars", level=2)
    doc.add_paragraph("{% for pillar in thesis_pillars %}")
    p = doc.add_paragraph()
    run = p.add_run("{{ pillar.title }}")
    run.bold = True
    doc.add_paragraph("{{ pillar.description }}")
    doc.add_paragraph("{% endfor %}")


def _add_company_overview(doc):
    doc.add_heading("Company Overview", level=1)
    doc.add_paragraph("{{ company_description }}")
    doc.add_paragraph("{{ segment_chart }}")


def _add_financial_analysis(doc):
    doc.add_heading("Financial Analysis", level=1)
    doc.add_paragraph("{{ revenue_chart }}")
    doc.add_paragraph("{{ margin_chart }}")
    # Table built dynamically by renderer as Subdoc
    doc.add_paragraph("{{ financials_table }}")


def _add_guidance_track_record(doc):
    doc.add_paragraph("{% if has_guidance %}")
    doc.add_heading("Guidance Track Record", level=1)
    doc.add_paragraph("{{ guidance_summary }}")
    # Table built dynamically by renderer as Subdoc
    doc.add_paragraph("{{ guidance_table }}")
    doc.add_paragraph("{% endif %}")


def _add_scenario_analysis(doc):
    doc.add_heading("Scenario Analysis", level=1)
    doc.add_paragraph("{{ scenario_chart }}")

    doc.add_heading("Bull Case", level=2)
    doc.add_paragraph("Probability: {{ bull_probability }}")
    doc.add_paragraph("Price Target: {{ bull_price_target }}")
    doc.add_paragraph("{{ bull_description }}")

    doc.add_heading("Base Case", level=2)
    doc.add_paragraph("Probability: {{ base_probability }}")
    doc.add_paragraph("Price Target: {{ base_price_target }}")
    doc.add_paragraph("{{ base_description }}")

    doc.add_heading("Bear Case", level=2)
    doc.add_paragraph("Probability: {{ bear_probability }}")
    doc.add_paragraph("Price Target: {{ bear_price_target }}")
    doc.add_paragraph("{{ bear_description }}")


def _add_capital_allocation(doc):
    doc.add_heading("Capital Allocation", level=1)
    doc.add_paragraph("{{ capital_allocation_summary }}")


def _add_valuation(doc):
    doc.add_heading("Valuation", level=1)

    doc.add_paragraph("{% if has_dcf %}")
    doc.add_heading("DCF Analysis", level=2)
    doc.add_paragraph("{{ dcf_summary }}")
    doc.add_paragraph("{{ dcf_sensitivity_chart }}")
    doc.add_paragraph("{% endif %}")

    doc.add_paragraph("{% if has_comps %}")
    doc.add_heading("Comparable Companies", level=2)
    # Table built dynamically by renderer as Subdoc
    doc.add_paragraph("{{ comps_table }}")
    doc.add_paragraph("{% endif %}")


def _add_risks(doc):
    doc.add_heading("Key Risks", level=1)
    # Table built dynamically by renderer as Subdoc
    doc.add_paragraph("{{ risks_table }}")


def _add_appendix(doc):
    doc.add_heading("Appendix", level=1)
    doc.add_paragraph("{{ appendix_content }}")
    doc.add_paragraph("")
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("Data sourced from Daloopa")
    run.font.size = Pt(9)
    run.font.color.rgb = MEDIUM_GRAY
    run.font.italic = True


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    doc = Document()
    _configure_styles(doc)
    _set_page_margins(doc)

    _add_cover_page(doc)
    _add_executive_summary(doc)
    _add_investment_thesis(doc)
    _add_company_overview(doc)
    _add_financial_analysis(doc)
    _add_guidance_track_record(doc)
    _add_scenario_analysis(doc)
    _add_capital_allocation(doc)
    _add_valuation(doc)
    _add_risks(doc)
    _add_appendix(doc)

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
